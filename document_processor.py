import fitz
import docx
import email
from email import policy
from typing import List
import mimetypes


def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF file."""
    doc = fitz.open(pdf_path)
    full_text = []
    for page in doc:
        text = page.get_text()
        if not text.strip():
            text = page.get_text("blocks")
        full_text.append(
            text if isinstance(text, str) else "\n".join([str(b) for b in text])
        )
    return "\n".join(full_text)


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from DOCX file."""
    doc = docx.Document(docx_path)
    full_text = []

    # Extract paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(" | ".join(row_text))

    return "\n".join(full_text)


def extract_text_from_email(email_path: str) -> str:
    """Extract text from email file (.eml)."""
    with open(email_path, "rb") as f:
        msg = email.message_from_bytes(f.read(), policy=policy.default)

    full_text = []

    # Extract headers
    if msg["Subject"]:
        full_text.append(f"Subject: {msg['Subject']}")
    if msg["From"]:
        full_text.append(f"From: {msg['From']}")
    if msg["To"]:
        full_text.append(f"To: {msg['To']}")
    if msg["Date"]:
        full_text.append(f"Date: {msg['Date']}")

    full_text.append("")  # Empty line separator

    # Extract body
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == "text/plain":
                body = part.get_content()
                if body:
                    full_text.append(body)
            elif content_type == "text/html":
                # Basic HTML stripping (you might want to use BeautifulSoup for better parsing)
                html_body = part.get_content()
                if html_body:
                    # Simple HTML tag removal
                    import re

                    clean_text = re.sub(r"<[^>]+>", "", html_body)
                    full_text.append(clean_text)
    else:
        body = msg.get_content()
        if body:
            full_text.append(body)

    return "\n".join(full_text)


def extract_text_from_document(file_path: str) -> str:
    """Extract text from PDF, DOCX, or email files."""
    # Determine file type by extension
    file_ext = file_path.lower().split(".")[-1]

    if file_ext == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_ext in ["docx", "doc"]:
        return extract_text_from_docx(file_path)
    elif file_ext in ["eml", "msg", "email"]:
        return extract_text_from_email(file_path)
    else:
        # Try to detect by content type
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type:
            if "pdf" in mime_type:
                return extract_text_from_pdf(file_path)
            elif "word" in mime_type or "document" in mime_type:
                return extract_text_from_docx(file_path)
            elif "message" in mime_type or "email" in mime_type:
                return extract_text_from_email(file_path)

        raise ValueError(
            f"Unsupported file format: {file_ext}. Supported formats: PDF, DOCX, EML"
        )


def chunk_text(text: str, chunk_size=500, overlap=50) -> List[str]:
    """Split text into chunks."""
    lines = text.split("\n")
    chunks = []
    chunk = ""
    for line in lines:
        if len(chunk) + len(line) < chunk_size:
            chunk += line.strip() + " "
        else:
            chunks.append(chunk.strip())
            chunk = line.strip() + " "
    if chunk:
        chunks.append(chunk.strip())
    return chunks
